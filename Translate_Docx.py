import docx
import asyncio
from googletrans import Translator
from docx import Document
from pypdf import PdfReader

def bossing_please_translate(dir_input, dir_output,src_lang,dest_lang):
    doc = Document(dir_input)
    allText=[]
    for para in doc.paragraphs:
        translator = Translator()
        new = asyncio.run(translator.translate(para.text,src=src_lang,dest=dest_lang))
        allText.append(new.text)
    doc = docx.Document()
    for para in allText:
        doc.add_paragraph(para)
    doc.save(dir_output)
    
def pdf_to_docx(in_path, out_path):
    allValue =[]
    reader = PdfReader(in_path)
    i = 0
    limit = (len(reader.pages))              
    while i < limit:
        page = reader.pages[i]
        value = (page.extract_text())
        allValue.append(value)
        i+=1   
    doc = docx.Document()
    for para in allValue:
        doc.add_paragraph(para)
    doc.save(out_path)

def main(pdf_in,pdf_out,src,dest):
    pdf_to_docx(pdf_in,pdf_out)
    bossing_please_translate(pdf_out,pdf_out,src,dest)

if __name__ == "__main__":
    en_lang = "en"
    pt_lang = "pt"
    pdf_path = "C:\\Debranding_PDF\\English Document Sample.pdf"
    sample_path = "C:\\Debranding_PDF\\English Document Sample_pt.docx"
    en_path = "C:\\Debranding_PDF\\English Document Sample_en.docx"
    main(pdf_path, sample_path, en_lang, pt_lang)
    main(pdf_path, en_path, en_lang, en_lang)
